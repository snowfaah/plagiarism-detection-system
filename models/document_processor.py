"""
document_processor.py - Handles PDF, DOCX, TXT extraction and preprocessing.
"""
import os
import re
import chardet
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(filepath: str) -> dict:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        
        if doc.is_encrypted:
            return {"error": "PDF is password-protected. Please provide an unlocked document.", "pages": 0}
        
        pages_text = []
        total_text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            pages_text.append({
                "page": page_num + 1,
                "text": text.strip(),
                "word_count": len(text.split())
            })
            total_text += text + "\n"
        
        doc.close()
        
        if not total_text.strip():
            return {"error": "PDF appears to be image-based or has no extractable text. Try OCR preprocessing.", "pages": len(pages_text)}
        
        return {
            "text": clean_text(total_text),
            "pages": len(pages_text),
            "pages_detail": pages_text,
            "word_count": len(total_text.split()),
            "char_count": len(total_text)
        }
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return {"error": f"Failed to process PDF: {str(e)}", "pages": 0}


def extract_text_from_docx(filepath: str) -> dict:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        
        paragraphs = []
        full_text = ""
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
                full_text += para.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text += cell.text + " "
        
        if not full_text.strip():
            return {"error": "DOCX appears to be empty or has no readable content.", "pages": 0}
        
        # Estimate pages (roughly 300 words/page)
        word_count = len(full_text.split())
        estimated_pages = max(1, word_count // 300)
        
        return {
            "text": clean_text(full_text),
            "pages": estimated_pages,
            "paragraphs": len(paragraphs),
            "word_count": word_count,
            "char_count": len(full_text)
        }
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return {"error": f"Failed to process DOCX: {str(e)}", "pages": 0}


def extract_text_from_txt(filepath: str) -> dict:
    """Extract text from TXT/HTML files with encoding detection."""
    try:
        with open(filepath, 'rb') as f:
            raw_bytes = f.read()
        
        detected = chardet.detect(raw_bytes)
        encoding = detected.get('encoding', 'utf-8') or 'utf-8'
        
        text = raw_bytes.decode(encoding, errors='replace')
        
        # Strip HTML if needed
        if filepath.endswith('.html') or filepath.endswith('.htm'):
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(text, 'html.parser')
            text = soup.get_text(separator='\n')
        
        if not text.strip():
            return {"error": "Text file appears to be empty.", "pages": 0}
        
        word_count = len(text.split())
        estimated_pages = max(1, word_count // 300)
        
        return {
            "text": clean_text(text),
            "pages": estimated_pages,
            "word_count": word_count,
            "char_count": len(text),
            "encoding": encoding
        }
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        return {"error": f"Failed to process text file: {str(e)}", "pages": 0}


def extract_text(filepath: str) -> dict:
    """Route file to correct extractor based on extension."""
    ext = os.path.splitext(filepath)[1].lower()
    
    extractors = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.doc': extract_text_from_docx,
        '.txt': extract_text_from_txt,
        '.html': extract_text_from_txt,
        '.htm': extract_text_from_txt,
    }
    
    extractor = extractors.get(ext)
    if not extractor:
        return {"error": f"Unsupported file format: {ext}. Supported: PDF, DOCX, TXT, HTML"}
    
    result = extractor(filepath)
    result['format'] = ext.lstrip('.')
    return result


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Remove non-printable characters except newlines
    text = re.sub(r'[^\x20-\x7E\n\u00A0-\uFFFF]', ' ', text)
    
    return text.strip()


def split_into_chunks(text: str, chunk_size: int = 200, overlap: int = 50) -> list:
    """Split text into overlapping chunks for analysis."""
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk = ' '.join(words[i:i + chunk_size])
        if len(chunk.split()) >= 20:  # Minimum chunk size
            chunks.append({
                "index": len(chunks),
                "text": chunk,
                "word_start": i,
                "word_end": min(i + chunk_size, len(words))
            })
    
    return chunks


def detect_language(text: str) -> str:
    """Detect document language."""
    try:
        from langdetect import detect
        sample = text[:2000]  # Use first 2000 chars
        return detect(sample)
    except Exception:
        return "en"


def validate_file(filepath: str, max_size_mb: int = 10) -> dict:
    """Validate file size and format."""
    if not os.path.exists(filepath):
        return {"valid": False, "error": "File not found"}
    
    size_bytes = os.path.getsize(filepath)
    size_mb = size_bytes / (1024 * 1024)
    
    if size_mb > max_size_mb:
        return {"valid": False, "error": f"File too large: {size_mb:.1f}MB. Maximum: {max_size_mb}MB"}
    
    ext = os.path.splitext(filepath)[1].lower()
    allowed = {'.pdf', '.docx', '.doc', '.txt', '.html', '.htm'}
    
    if ext not in allowed:
        return {"valid": False, "error": f"Unsupported format: {ext}"}
    
    return {"valid": True, "size_mb": round(size_mb, 2)}